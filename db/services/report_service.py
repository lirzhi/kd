from copy import deepcopy
import logging
import os
import json

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from db.dbutils import singleton
from db.dbutils.redis_conn import RedisDB
from db.ectd_model import EctdSectionModel
from db.services.file_service import FileService
from utils.file_util import ensure_dir_exists
from mutil_agents.memory.specific_review_state import SpecificReviewState
from mutil_agents.agent import specific_report_generationAgent_graph
from utils.common_util import produce_handle_info

ReportInfo = {
        "status": 0,
        "message": "",
        "data": None
}

REPORT_DIR = "data/reports/"  # Set the target folder for report generation
ensure_dir_exists(REPORT_DIR)
@singleton
class ReportService:
    def __init__(self):
        self.redis_conn = RedisDB()
        self._report_status = {}  # 用于存储报告生成状态

    def get_report_status(self, doc_id):
        """获取报告生成状态"""
        status = self._report_status.get(doc_id, {
            'status': 'pending',  # pending, processing, completed, failed
            'progress': 0,
            'message': ''
        })
        return status

    def update_report_status(self, doc_id, status, progress=0, message=''):
        """更新报告生成状态"""
        self._report_status[doc_id] = {
            'status': status,
            'progress': progress,
            'message': message
        }

    def ectd_judge(self, doc_id):
        resp_info = deepcopy(ReportInfo)
        # 根据eCTD的doc_id生成最终报告
        # 获取doc_id对应的文件信息
        
        file_info = FileService().get_file_by_id(doc_id)
        if file_info is None:
            resp_info["message"] = "文件不存在" 
            return resp_info
        
        if file_info.classification != "eCTD":
            resp_info["message"] = "文件类型不是eCTD"
            return resp_info
        
        if file_info.is_chunked != 1:
            resp_info["message"] = "eCTD未解析,解析后在试"
            return resp_info   
        # 所有检查通过，标记为成功
        resp_info["status"] = 1
        resp_info["message"] = "eCTD文件验证通过"
        resp_info["data"] = file_info  # 可选：返回文件信息供后续使用
        return resp_info 

    def generate_report_by_section(self, doc_id, section_id, section_name, content):
        # 按章节生成内容,生成完成后将内容存入数据库，并返回生成的内容
        produce_handle_info({"task": "【当前处理进度】", "data":f"章节ID：{section_id}，章节名称：{section_name}"})
        review_state = SpecificReviewState()
        review_state["content"] = content
        review_state["content_section"] = section_id
        review_state["content_section_name"] = section_name
        review_require_list = self.redis_conn.get(f"principle+{section_id}")
        if review_require_list is None:
            logging.warning(f"{section_id}章节要求不存在")
            review_require_list = []
        review_state["review_require_list"] = review_require_list
        result = specific_report_generationAgent_graph.invoke(review_state)
        flag = self.redis_conn.set(f"review_content+{doc_id}+{section_id}", result["final_report_content"], None)
        if flag:
            print("true")
        else:
            print("false")
            print(result["final_report_content"])
        return result
        

    def create_final_report(self, doc_id):
        """异步生成完整报告"""
        resp_info = deepcopy(ReportInfo)
        
        # 检查文件是否有效
        judge_info = self.ectd_judge(doc_id)
        if judge_info["status"] == 0:
            self.update_report_status(doc_id, 'failed', message=judge_info["message"])
            return judge_info

        # 更新状态为处理中
        self.update_report_status(doc_id, 'processing', progress=0)
        
        try:
            report = deepcopy(EctdSectionModel)
            total_sections = len(report)
            completed_sections = 0

            # 遍历所有章节
            for section in report:
                if len(section["children_sections"]) == 0:
                    # 处理无子章节的情况
                    section_id = section["section_id"]
                    section_name = section["section_name"]
                    section_content = self.redis_conn.get(f"section_content+{doc_id}+{section_id}")
                    
                    if section_content:
                        section_report = self.generate_report_by_section(doc_id, section_id, section_name, section_content)
                        section["report"] = section_report["final_report"]
                    
                    completed_sections += 1
                    progress = int((completed_sections / total_sections) * 100)
                    self.update_report_status(doc_id, 'processing', progress=progress)
                else:
                    # 处理有子章节的情况
                    for child_section in section["children_sections"]:
                        child_section_id = child_section["section_id"]
                        child_section_name = child_section["section_name"]
                        child_section_content = self.redis_conn.get(f"section_content+{doc_id}+{child_section_id}")
                        
                        if child_section_content:
                            content = self.generate_report_by_section(doc_id, child_section_id, child_section_name, child_section_content)
                            child_section["report"] = content["final_report"]
                        
                        completed_sections += 1
                        progress = int((completed_sections / total_sections) * 100)
                        self.update_report_status(doc_id, 'processing', progress=progress)

            # 更新状态为完成
            self.update_report_status(doc_id, 'completed', progress=100)
            resp_info["status"] = 1
            resp_info["message"] = "报告生成成功"
            resp_info["data"] = report
            return resp_info

        except Exception as e:
            error_msg = f"生成报告失败: {str(e)}"
            self.update_report_status(doc_id, 'failed', message=error_msg)
            resp_info["status"] = 0
            resp_info["message"] = error_msg
            return resp_info

    def add_section_to_doc(self, doc_id, doc, section, level=0):
        """递归添加章节到文档"""
        # 添加章节标题
        doc.add_heading(f"{section['section_id']} {section['section_name']}", level=level)
        
        # 添加章节正文内容
        if self.redis_conn.exist(f"review_content+{doc_id}+{section['section_id']}"):
            content = self.redis_conn.get(f"review_content+{doc_id}+{section['section_id']}")
            if content:
                doc.add_paragraph(content)
        
        # 如果有子章节，递归添加
        for child in section.get("children_sections", []):
            self.add_section_to_doc(doc_id, doc, child, level=level + 1)

    def export_report(self, doc_id, export_type='word'):
        resp_info = deepcopy(ReportInfo)
        try:
            # 检查文件是否有效
            judge_info = self.ectd_judge(doc_id)
            if judge_info["status"] == 0:
                return judge_info
            
            if len(self.redis_conn.keys(f"review_content+{doc_id}*")) == 0:
                resp_info["message"] = "还没有章节生成报告内容"
                resp_info["status"] = 0
                return resp_info
            
            # 获取文件信息
            file_info = FileService().get_file_by_id(doc_id)
            if not file_info:
                resp_info["message"] = "文件不存在"
                resp_info["status"] = 0
                return resp_info
            
            # 创建报告目录
            os.makedirs(REPORT_DIR, exist_ok=True)
            
            # 生成Word文档
            report = deepcopy(EctdSectionModel)
            doc = Document()
            
            # 设置文档默认字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.styles['Normal'].font.size = Pt(12)
            
            # 添加文档标题
            doc.add_heading(f"{file_info.file_name}审评报告", 0)
            
            # 遍历章节并添加到文档
            for section in report:
                self.add_section_to_doc(doc_id, doc, section)
            
            # 生成文件名（使用原始文件名作为基础）
            base_name = os.path.splitext(file_info.file_name)[0]
            word_filename = f"{base_name}_报告.docx"
            pdf_filename = f"{base_name}_报告.pdf"
            
            # 保存Word文档
            word_path = os.path.join(REPORT_DIR, word_filename)
            doc.save(word_path)
            
            if not os.path.exists(word_path):
                raise Exception("Word文件生成失败")
            
            if export_type == 'pdf':
                pdf_path = os.path.join(REPORT_DIR, pdf_filename)
                try:
                    import docx2pdf
                    # 转换文档
                    docx2pdf.convert(word_path, pdf_path)
                    
                    # 检查PDF文件是否生成成功
                    if not os.path.exists(pdf_path):
                        raise Exception("PDF文件生成失败")
                    
                    if os.path.getsize(pdf_path) == 0:
                        raise Exception("生成的PDF文件为空")
                    
                    resp_info["message"] = "PDF报告生成成功"
                    resp_info["status"] = 1
                    resp_info["data"] = pdf_path
                    return resp_info
                    
                except Exception as e:
                    logging.error(f"PDF转换失败: {str(e)}")
                    # 如果PDF转换失败，至少返回Word文件
                    resp_info["message"] = f"PDF转换失败，已生成Word文件: {str(e)}"
                    resp_info["status"] = 1
                    resp_info["data"] = word_path
                    return resp_info
            
            resp_info["message"] = "Word报告生成成功"
            resp_info["status"] = 1
            resp_info["data"] = word_path
            return resp_info
            
        except Exception as e:
            logging.error(f"报告生成失败: {str(e)}")
            resp_info["message"] = f"报告生成失败: {str(e)}"
            resp_info["status"] = 0
            resp_info["data"] = None
            return resp_info

    def delete_report(self, doc_id):
        # Logic to delete a report from the database
        if self.redis_conn.exist(f"review_content+{doc_id}*"):
            self.redis_conn.delete(f"review_content+{doc_id}*")
        if os.path.exists(f"{REPORT_DIR}{doc_id}_report.docx"):
            os.remove(f"{REPORT_DIR}{doc_id}_report.docx")
            return True
        else:
            logging.warning(f"Report file {doc_id}_report.docx does not exist.")
            return False
        
    def get_report_content(self, doc_id):
        """获取完整报告内容"""
        resp_info = deepcopy(ReportInfo)
        logging.info(f"开始获取报告内容，doc_id: {doc_id}")
        
        # 检查文件是否有效
        judge_info = self.ectd_judge(doc_id)
        if judge_info["status"] == 0:
            logging.error(f"文件验证失败: {judge_info['message']}")
            return judge_info
        
        # 检查redis中是否已存在完整报告
        report_key = f"report+{doc_id}"
        existing_report = self.redis_conn.get(report_key)
        
        if existing_report:
            logging.info(f"从Redis获取到已存在的报告内容，key: {report_key}")
            resp_info["status"] = 1
            resp_info["data"] = existing_report
            return resp_info
        
        # 检查是否有章节报告内容
        review_keys = self.redis_conn.keys(f"review_content+{doc_id}*")
        if not review_keys:
            logging.error(f"未找到任何章节报告内容，doc_id: {doc_id}")
            resp_info["status"] = 0
            resp_info["message"] = "未找到任何章节报告内容"
            return resp_info
        
        # 获取章节模型
        report = deepcopy(EctdSectionModel)
        content = []
        
        def add_section_content(sections, level=0):
            """递归添加章节内容"""
            for section in sections:
                # 添加章节标题
                title = f"{section['section_id']} {section['section_name']}"
                content.append(f"{'#' * (level + 1)} {title}\n")
                
                # 添加章节正文内容
                section_key = f"review_content+{doc_id}+{section['section_id']}"
                section_content = self.redis_conn.get(section_key)
                
                if section_content:
                    try:
                        # 如果内容是JSON字符串，解析它
                        if isinstance(section_content, str):
                            section_content = json.loads(section_content)
                        content.append(f"{section_content}\n")
                    except json.JSONDecodeError:
                        content.append(f"{section_content}\n")
                
                # 如果有子章节，递归添加
                if section.get("children_sections"):
                    add_section_content(section["children_sections"], level + 1)
        
        try:
            # 生成内容
            add_section_content(report)
            report_content = '\n'.join(content)
            
            if not report_content.strip():
                logging.error("生成的报告内容为空")
                resp_info["status"] = 0
                resp_info["message"] = "生成的报告内容为空"
                return resp_info
            
            # 存储到redis
            success = self.redis_conn.set(report_key, report_content)
            if not success:
                logging.error(f"存储报告内容到Redis失败，key: {report_key}")
                resp_info["status"] = 0
                resp_info["message"] = "存储报告内容失败"
                return resp_info
            
            logging.info(f"成功生成并存储报告内容，key: {report_key}")
            resp_info["status"] = 1
            resp_info["data"] = report_content
            return resp_info
            
        except Exception as e:
            logging.error(f"生成报告内容时发生错误: {str(e)}")
            resp_info["status"] = 0
            resp_info["message"] = f"生成报告内容失败: {str(e)}"
            return resp_info
        