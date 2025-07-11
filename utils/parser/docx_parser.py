from docx import Document
import os

from db.dbutils import singleton
from utils.file_util import ensure_dir_exists

DOCX_IMAGE_PATH = 'data/parser/docx/images/'
DOCX_TABLE_PATH = "data/parser/docx/tables/"
@singleton
class DocxParser:
    def __init__(self):
        ensure_dir_exists(DOCX_IMAGE_PATH)
        ensure_dir_exists(DOCX_TABLE_PATH)
   
    def read_word_in_chunks(self, file_path, chunk_size):
        doc = Document(file_path)
        chunks = []
        current_chunk = {
            'text': '',
            'tables': [],
            'images': [],
            'start_pos': None
        }
        current_chunk_index = 0

        for i, paragraph in enumerate(doc.paragraphs):
            # 当前段落的文本
            paragraph_text = paragraph.text.strip()
            if paragraph_text:  # 忽略空段落
                # 更新当前chunk的起始位置
                if current_chunk['start_pos'] is None:
                    current_chunk['start_pos'] = i
                    current_chunk["chunk_id"] = i + 1
                # 使用spacy进行句子分割
                doc_text = paragraph_text.split('。')
                for sent in doc_text:
                    # 获取句子的文本
                    sentence_text = sent.strip() + "。"
                    # 检查是否达到chunk大小限制
                    if len(current_chunk['text']) + len(sentence_text) > chunk_size:
                        # 保存当前chunk并开始新的chunk
                        chunks.append(current_chunk)
                        current_chunk = {
                            'text': sentence_text,
                            'tables': [],
                            'image_paths': [],
                            'start_pos': i,
                            "chunk_id": i + 1
                        }
                        current_chunk_index += 1
                    else:
                        # 添加句子文本到当前chunk
                        current_chunk['text'] += sentence_text + '\n'

        # 添加最后一个chunk（如果有）
        if current_chunk['text']:
            chunks.append(current_chunk)

        # 处理表格和图片，分别保存
        if not chunks:
            return []  # 没有内容直接返回空

        image_dir = os.path.join(DOCX_IMAGE_PATH)
        if not os.path.exists(image_dir):
            os.makedirs(image_dir)

        table_dir = os.path.join(DOCX_TABLE_PATH)
        if not os.path.exists(table_dir):
            os.makedirs(table_dir)

        # 遍历文档中的表格和图片，将它们添加到相应的chunk中
        for i, table in enumerate(doc.tables):
            # 只分配给有 start_pos 的 chunk
            valid_chunks = [(chunk['start_pos'], idx) for idx, chunk in enumerate(chunks) if chunk.get('start_pos') is not None]
            if not valid_chunks:
                continue
            table_chunk_index = min(valid_chunks)[1]
            chunks[table_chunk_index]['tables'].append({
                'index': i,
                'text': " ".join([cell.text for row in table.rows for cell in row.cells])
            })

        for i, inline_shape in enumerate(doc.inline_shapes):
            if inline_shape.type == 3:  # 图片类型
                image = inline_shape.image
                image_path = os.path.join(image_dir, f'image_{i}.png')
                with open(image_path, 'wb') as img_file:
                    img_file.write(image)
                valid_chunks = [(chunk['start_pos'], idx) for idx, chunk in enumerate(chunks) if chunk.get('start_pos') is not None]
                if not valid_chunks:
                    continue
                image_chunk_index = min(valid_chunks)[1]
                chunks[image_chunk_index]['image_paths'].append(image_path)
        return chunks
   
    def __call__(self, path):
        return self.read_word_in_chunks(path, 512)